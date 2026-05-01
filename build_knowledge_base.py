import os
import pickle
import numpy as np
from pathlib import Path
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss

def read_docx(path):
    from docx import Document
    doc = Document(path)
    content = []
    # Lire les paragraphes
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(p.text.strip())
    # Lire les tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() 
                for cell in row.cells 
                if cell.text.strip()
            )
            if row_text:
                content.append(row_text)
    return "\n".join(content)

def split_chunks(text, size=800, overlap=150):
    chunks, start = [], 0
    while start < len(text):
        chunk = text[start:start+size]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += size - overlap
    return chunks

docs_path = Path("docs")
all_chunks, all_sources = [], []

print("📂 Lecture des documents...")
for f in docs_path.iterdir():
    if f.suffix.lower() == ".docx":
        try:
            text = read_docx(str(f))
            chunks = split_chunks(text)
            all_chunks.extend(chunks)
            all_sources.extend([f.name] * len(chunks))
            print(f"   ✅ {f.name} → {len(chunks)} chunks")
        except Exception as e:
            print(f"   ❌ {f.name} : {e}")

print(f"\n✅ Total : {len(all_chunks)} chunks")
print("⏳ Calcul embeddings...")

model = SentenceTransformer("all-MiniLM-L6-v2")
embeddings = model.encode(all_chunks, batch_size=8, show_progress_bar=True)

with open("knowledge_base.pkl", "wb") as f:
    pickle.dump({"chunks": all_chunks, "sources": all_sources, "embeddings": embeddings}, f)

print("✅ Base prête ! → Lancez : python app.py")
